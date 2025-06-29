import streamlit as st
import pandas as pd
import joblib
import numpy as np
from datetime import datetime
import os
import matplotlib
import matplotlib.pyplot as plt
from docx import Document

# Set Matplotlib backend to Agg for Streamlit compatibility
matplotlib.use('Agg')

# Set page config
st.set_page_config(page_title="IntelliHealth: AI-Based Disease Predictor", layout="wide")

# Custom CSS
st.markdown(
    """
    <style>
      .stApp {
        background-image: url("https://www.boffinaccess.com/public/frontend/recently-published-articles/a-case-of-post-COVID-19.jpg");
        background-size: cover;
        background-repeat: no-repeat;
        background-attachment: fixed;
        background-position: center;
        height: 100vh;
      }
      .main::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background: rgba(255, 255, 255, 0.85);
        z-index: 1;
      }
      .main .block-container {
        padding-left: 1rem;
        padding-right: 1rem;
        padding-top: 0.5rem;
        max-width: 95%;
        position: relative;
      }
      .stButton > button {
        background-color: #000080;
        color: #FFFFFF;
        border-radius: 5px;
        padding: 10px 20px;
        font-size: 16px;
        border: none;
        margin: 5px;
        min-width: 120px;
      }
      .stButton > button:hover {
        background-color: #0000A0;
        color: #FFFFFF;
        cursor: pointer;
      }
      .stButton > button:active {
        background-color: #0000A0;
        color: #FFFFFF !important;
        cursor: pointer;
      }
      .stButton > button:disabled {
        background-color: #4B4B4B;
        color: #FFFFFF;
        cursor: not-allowed;
      }
      .stButton > button:disabled:hover {
        background-color: #D3D3D3;
        color: #FFFFFF;
        cursor: not-allowed;
      }
      .button-container {
        display: flex;
        justify-content: center;
        align-items: center;
        gap: 10px;
        margin-top: 20px;
        margin-bottom: 20px;
      }
      .button-container .stButton:first-child > button {
        margin-left: 15px;
      }
       .predict-button {
        margin-left: 50px !important;
        margin-right: 50px !important;
      }
      .prediction-container {
    background-color: #fefefe;
    padding: 20px;
    border-radius: 10px;
    box-shadow: 0 4px 8px rgba(0,0,0,0.2);
    margin-top: 20px;
    width: 50%;
    margin-left: auto;
    margin-right: auto;
}
      .prediction-row {
        display: flex;
        flex-wrap: wrap;
        justify-content: space-between;
        gap: 15px;
        margin-bottom: 20px;
      }
      .prediction-card {
        flex: 1;
        min-width: 250px;
        height: 400px;
        background-color: #e6f0ff;
        padding: 20px;
        border-radius: 12px;
        border: 2px solid #1e90ff;
        box-shadow: 0 4px 12px rgba(0,0,0,0.15);
        text-align: center;
        transition: transform 0.2s;
        display: flex;
        flex-direction: column;
        justify-content: space-between;
        box-sizing: border-box;
      }
      .prediction-card:hover {
        transform: translateY(-5px);
      }
      .prediction-card h3 {
        color: #000080;
        margin: 0 0 8px 0;
        font-size: 24px;
        font-weight: 600;
      }
      .prediction-card p {
        margin: 8px 0;
        font-size: 18px;
        color: #333;
      }
      .prediction-card p strong {
        color: #000080;
      }
      .recommendation {
        background-color: #f0f8ff;
        padding: 10px;
        margin-top: 15px;
        border-radius: 8px;
        border-left: 5px solid #1e90ff;
        box-shadow: 0 2px 6px rgba(0,0,0,0.1);
        font-size: 16px;
        color: #333;
        line-height: 1.6;
        text-align: left;
        flex-grow: 1;
        overflow-y: auto;
        box-sizing: border-box;
      }
      .recommendation strong {
        color: #000080;
        font-weight: 600;
      }
      .stNumberInput input {
        background-color: #f5f5f5;
        border: 1px solid #ccc;
        border-radius: 4px;
        padding: 8px;
      }
      .stSelectbox select {
        background-color: #f5f5f5;
        border: 1px solid #ccc;
        border-radius: 4px;
        padding: 8px;
      }
      .stNumberInput label, .stSelectbox label {
        font-size: 18px;
        color: #333;
      }
      .stMarkdown p {
        font-size: 18px;
        color: #333;
      }
      .header-container {
        display: flex;
        justify-content: space-between;
        align-items: center;
        background-color: #ADD8E6;
        padding: 10px;
        border-radius: 5px;
        margin-top: 0.2rem;
        position: relative;
      }
      .header-title {
        color: #FFFFFF;
        font-size: 16px;
        margin: 0;
      }
      .header-right {
        display: flex;
        align-items: center;
        gap: 10px;
      }
      .profile-icon {
        font-size: 24px;
        color: #00008B;
        margin-right: 10px;
        background-color: #FFFFFF;
        border-radius: 50%;
        padding: 5px;
        cursor: pointer;
      }
      .about-button {
        background-color: #000080;
        color: #FFFFFF;
        border-radius: 5px;
        padding: 2px 7px;
        font-size: 24px;
        border: none;
        cursor: pointer;
      }
      .about-button:hover {
        background-color: #0000A0;
      }
      [data-testid="stSidebar"] {
        color: #FFFFFF;
      }
      [data-testid="stSidebarNav"] a {
        color: #FFFFFF !important;
        font-size: 16px;
      }
      [data-testid="stSidebarNav"] a:hover {
        background-color: #0000A0;
        color: #FFFFFF !important;
      }
      .center-button-container {
        display: flex;
        justify-content: center;
        margin-top: 20px;
      }
      .question-container {
        background-color: #fefefe;
        padding: 20px;
        border-radius: 10px;
        border: 2px solid #1e90ff;
        box-shadow: 0 4px 8px rgba(0,0,0,0.2);
        margin: 20px 0;
        display: flex;
        flex-direction: column;
        gap: 15px;
      }
      .chart-container {
        background-color: white;
        padding: 20px;
        border-radius: 10px;
        box-shadow: 0 4px 8px rgba(0,0,0,0.1);
        margin-top: 20px;
      }
     .chart-header {
    margin-top: 20px;
    background-color: #ffffff;
    padding: 8px;
    border-radius: 5px;
    text-align: center;
    margin-bottom: 15px;
    width: 50%;
    margin-left: auto;
    margin-right: auto;
    border: 1px solid #ADD8E6;
}
.chart-header h3 {
    color: #000080;
    margin: 0;
    font-size: 20px;
    font-weight: 600;
}
    </style>
    """,
    unsafe_allow_html=True
)

# Initialize session state for About page
if 'show_about' not in st.session_state:
    st.session_state.show_about = False

# Header with profile icon and About icon button
st.markdown(
    """
    <div class="header-container">
        <h1 class="header-title">
            <span style="color: #00008B;">🩺</span> IntelliHealth: AI-Based Disease Predictor
        </h1>
        <div class="header-right">
            <button class="about-button" onclick="document.getElementById('about_button').click()">ℹ️</button>
            <span class="profile-icon">👤</span>
        </div>
    </div>
    """,
    unsafe_allow_html=True
)

# Load models and encoders
try:
    diabetes_model = joblib.load('best_model_Risk_Level_Diabetes.pkl')
    heart_disease_model = joblib.load('best_model_Risk_Level_Heart_Disease.pkl')
    hypertension_model = joblib.load('best_model_Risk_Level_Hypertension.pkl')
    encoders = joblib.load('label_encoders.pkl')
    feature_sets = {
        'Risk_Level_Diabetes': joblib.load('features_Risk_Level_Diabetes.pkl'),
        'Risk_Level_Hypertension': joblib.load('features_Risk_Level_Hypertension.pkl'),
        'Risk_Level_Heart_Disease': joblib.load('features_Risk_Level_Heart_Disease.pkl')
    }
except FileNotFoundError as e:
    st.error(f"Required file missing: {str(e)}. Ensure 'best_model_{{target}}.pkl', 'label_encoders.pkl', and 'features_{{target}}.pkl' are in the same directory.")
    st.stop()

# Define input mappings to match encoder classes
input_mappings = {
    'Gender': {'Male': 'Male', 'Female': 'Female'},
    'Smoking': {'No': 'No', 'Yes': 'Yes'},
    'Alcohol_Use': {'No': 'No', 'Yes': 'Yes'},
    'FruitVegFreq': {'Daily': 'Daily', 'Sometimes': 'Sometimes', 'Rarely': 'Rarely'},
    'MealType': {'Home-cooked': 'Home-cooked', 'Mix': 'Mix', 'Takeaways': 'Takeaways'},
    'SugaryIntake': {'Rarely': 'Rarely', 'Moderately': 'Moderately', 'Frequently': 'Frequently'},
    'ReadLabels': {'No': 'No', 'Sometimes': 'Sometimes', 'Yes': 'Yes'},
    'ExerciseDays': {'Rarely': 'Rarely', 'Moderately': 'Moderately', 'Regularly': 'Regularly'},
    'WalkCycle': {'No': 'No', 'Sometimes': 'Sometimes', 'Yes': 'Yes'},
    'DeskJob': {'No': 'No', 'Partially': 'Partially', 'Yes': 'Yes'},
    'PhysicalTired': {'No': 'No', 'Sometimes': 'Sometimes', 'Yes': 'Yes'},
    'Anxious': {'No': 'No', 'Sometimes': 'Sometimes', 'Yes': 'Yes'},
    'GoodSleep': {'No': 'No', 'Yes': 'Yes'},
    'MoodSwings': {'No': 'No', 'Occasionally': 'Occasionally', 'Yes': 'Yes'},
    'Relaxed': {'No': 'No', 'Sometimes': 'Sometimes', 'Yes': 'Yes'},
    'ThirstyFatigued': {'No': 'No', 'Sometimes': 'Sometimes', 'Yes': 'Yes'},
    'WeightChange': {'No': 'No', 'Not sure': 'Not sure', 'Yes': 'Yes'},
    'ChestPain': {'No': 'No', 'Sometimes': 'Sometimes', 'Yes': 'Yes'},
    'FrequentUrination': {'No': 'No', 'Sometimes': 'Sometimes', 'Yes': 'Yes'},
    'BlurredVision': {'No': 'No', 'Sometimes': 'Sometimes', 'Yes': 'Yes'},
    'FamDiabetes': {'No': 'No', 'Not sure': 'Not sure', 'Yes': 'Yes'},
    'FamHeartDisease': {'No': 'No', 'Not sure': 'Not sure', 'Yes': 'Yes'},
    'FamHypertension': {'No': 'No', 'Not sure': 'Not sure', 'Yes': 'Yes'},
    'FamCholesterol': {'No': 'No', 'Not sure': 'Not sure', 'Yes': 'Yes'},
    'FamStroke': {'No': 'No', 'Not sure': 'Not sure', 'Yes': 'Yes'},
    'FamObesity': {'No': 'No', 'Not sure': 'Not sure', 'Yes': 'Yes'},
    'FamHeartSurgery': {'No': 'No', 'Not sure': 'Not sure', 'Yes': 'Yes'},
    'FamChronic': {'No': 'No', 'Not sure': 'Not sure', 'Yes': 'Yes'},
    'FamKidney': {'No': 'No', 'Not sure': 'Not sure', 'Yes': 'Yes'},
    'FamEarlyDeath': {'No': 'No', 'Not sure': 'Not sure', 'Yes': 'Yes'}
}

# Initialize session state
if 'category_index' not in st.session_state:
    st.session_state.category_index = 0
if 'inputs' not in st.session_state:
    st.session_state.inputs = {
        'Age': 30,
        'Gender': 'Male',
        'BMI': 25.0,
        'Smoking': 'No',
        'Alcohol_Use': 'No',
        'Sleep_Hours': 7,
        'Systolic_BP': 120,
        'Diastolic_BP': 80,
        'Blood_Sugar': 100.0,
        'FruitVegFreq': 'Daily',
        'MealType': 'Home-cooked',
        'SugaryIntake': 'Rarely',
        'ReadLabels': 'No',
        'ExerciseDays': 'Rarely',
        'WalkCycle': 'No',
        'DeskJob': 'No',
        'PhysicalTired': 'No',
        'Anxious': 'No',
        'GoodSleep': 'No',
        'MoodSwings': 'No',
        'Relaxed': 'No',
        'ThirstyFatigued': 'No',
        'WeightChange': 'No',
        'ChestPain': 'No',
        'FrequentUrination': 'No',
        'BlurredVision': 'No',
        'FamDiabetes': 'No',
        'FamHeartDisease': 'No',
        'FamHypertension': 'No',
        'FamCholesterol': 'No',
        'FamStroke': 'No',
        'FamObesity': 'No',
        'FamHeartSurgery': 'No',
        'FamChronic': 'No',
        'FamKidney': 'No',
        'FamEarlyDeath': 'No'
    }
if 'show_results' not in st.session_state:
    st.session_state.show_results = False
if 'predictions' not in st.session_state:
    st.session_state.predictions = None
if 'diabetes_risk' not in st.session_state:
    st.session_state.diabetes_risk = None
if 'hypertension_risk' not in st.session_state:
    st.session_state.hypertension_risk = None
if 'heart_disease_risk' not in st.session_state:
    st.session_state.heart_disease_risk = None

# Function to save predictions to predictions.csv with specified format
def save_prediction(patient_id, disease, prediction):
    DATA_FILE = "predictions.csv"
    new_prediction = {
        "Prediction ID": f"P{int(datetime.now().timestamp())}",
        "Date": datetime.now(),
        "Patient ID": patient_id,
        "Disease": disease,
        "Prediction": prediction
    }
    try:
        if os.path.exists(DATA_FILE):
            df = pd.read_csv(DATA_FILE)
            df = pd.concat([df, pd.DataFrame([new_prediction])], ignore_index=True)
        else:
            df = pd.DataFrame([new_prediction])
        df.to_csv(DATA_FILE, index=False)
    except Exception as e:
        st.error(f"Error saving prediction to {DATA_FILE}: {str(e)}")

# Function to save inputs and predictions to a single Word document
def save_to_word(patient_id, inputs, diabetes_risk, hypertension_risk, heart_disease_risk):
    DOC_FILE = "all_user_data.docx"
    try:
        # Check if the file is accessible
        if os.path.exists(DOC_FILE):
            with open(DOC_FILE, 'a') as f:
                pass  # Just check if we can open it
            doc = Document(DOC_FILE)
        else:
            doc = Document()
            doc.add_heading('IntelliHealth User Data', level=0)

        # Add a separator and new user container section
        doc.add_paragraph("--------------------------------------------------")
        doc.add_heading(f"Patient ID: {patient_id}", level=1)

        # Add user inputs
        doc.add_heading('User Inputs', level=2)
        for key, value in inputs.items():
            doc.add_paragraph(f"- {key}: {value}")

        # Add prediction outputs
        doc.add_heading('Prediction Outputs', level=2)
        doc.add_paragraph(f"- Diabetes Risk: {diabetes_risk}")
        doc.add_paragraph(f"- Hypertension Risk: {hypertension_risk}")
        doc.add_paragraph(f"- Heart Disease Risk: {heart_disease_risk}")

        # Save the document
        doc.save(DOC_FILE)
        st.success(f"Data successfully appended to {DOC_FILE}")
    except PermissionError:
        st.error(f"Permission denied: Unable to save to {DOC_FILE}. Please ensure the file is not open in another program (e.g., Microsoft Word) and that you have write permissions to the directory. Close the file and try again.")
    except Exception as e:
        st.error(f"Error saving to {DOC_FILE}: {str(e)}. Please check file permissions or ensure the document is not open in another program.")

# About page content
if st.session_state.show_about:
    st.markdown(
        """
        <div style='background-color: #fefefe; padding: 20px; border-radius: 10px; box-shadow: 0 4px 8px rgba(0,0,0,0.2); margin-top: 20px;'>
            <h2 style='text-align: center; color: #000080;'>About IntelliHealth</h2>
            <p style='font-size: 18px; color: #333;'>
                <strong>IntelliHealth</strong> is an AI-based disease prediction platform designed to assess the risk of chronic conditions such as Diabetes, Hypertension, and Heart Disease. By leveraging advanced machine learning models, IntelliHealth analyzes user-provided health data to deliver accurate risk assessments and personalized health recommendations.
            </p>
            <p style='font-size: 18px; color: #333;'>
                <strong>Features:</strong><br>
                • Comprehensive input forms covering personal, lifestyle, and family health data.<br>
                • Real-time risk predictions with severity scores.<br>
                • Tailored recommendations for diet, activity, and medical follow-ups.<br>
                • Prediction history tracking via a companion app.
            </p>
            <p style='font-size: 18px; color: #333;'>
                <strong>Our Team:</strong><br>
                Developed by a team of data scientists and healthcare enthusiasts dedicated to empowering individuals with proactive health insights.
            </p>
            <p style='font-size: 18px; color: #333;'>
                <strong>Contact:</strong><br>
                For feedback or inquiries, reach out at <a href='mailto:support@intellihealth.ai'>support@intellihealth.ai</a>.
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
    if st.button("Back", key="back_button"):
        st.session_state.show_about = False
else:
    # Main content
    # Define categories
    categories = [
        "Personal Information",
        "Lifestyle & Daily Habits",
        "Health Symptoms",
        "Diet Type",
        "Physical Activity",
        "Stress Level",
        "Symptoms",
        "Family History"
    ]

    # Display current category
    current_category = categories[st.session_state.category_index]
    st.header(current_category)

    # Helper function for selectbox index
    def safe_selectbox_index(options, value, default_index=0):
        try:
            return options.index(value)
        except ValueError:
            return default_index

    # Validation function to check if all inputs in the current category are answered
    def validate_category_inputs(category, inputs):
        required_fields = {
            "Personal Information": ['Age', 'Gender', 'BMI'],
            "Lifestyle & Daily Habits": ['Smoking', 'Alcohol_Use', 'Sleep_Hours'],
            "Health Symptoms": ['Systolic_BP', 'Diastolic_BP', 'Blood_Sugar'],
            "Diet Type": ['FruitVegFreq', 'MealType', 'SugaryIntake', 'ReadLabels'],
            "Physical Activity": ['ExerciseDays', 'WalkCycle', 'DeskJob', 'PhysicalTired'],
            "Stress Level": ['Anxious', 'GoodSleep', 'MoodSwings', 'Relaxed'],
            "Symptoms": ['ThirstyFatigued', 'WeightChange', 'ChestPain', 'FrequentUrination', 'BlurredVision'],
            "Family History": ['FamDiabetes', 'FamHeartDisease', 'FamHypertension', 'FamCholesterol', 'FamStroke',
                              'FamObesity', 'FamHeartSurgery', 'FamChronic', 'FamKidney', 'FamEarlyDeath']
        }
        if category not in required_fields:
            st.warning(f"Unknown category: {category}. Validation skipped.")
            return False
        for field in required_fields[category]:
            value = inputs.get(field)
            if value is None:
                return False
            if isinstance(value, str) and not value:
                return False
            if isinstance(value, (int, float)) and (value < 0 or np.isnan(value)):
                return False
            # Check if categorical input is valid according to encoder
            if field in encoders:
                try:
                    encoders[field].transform([value])
                except ValueError:
                    return False
        return True

    # Helper function to validate and set selectbox input
    def validated_selectbox(label, options, current_value, field, key):
        selected = st.selectbox(label, options, index=safe_selectbox_index(options, current_value), key=key)
        if field in encoders:
            valid_options = list(encoders[field].classes_)
            if selected not in valid_options:
                st.error(f"Invalid value for {field}: {selected}. Please select one of {valid_options}.")
                default_value = valid_options[0] if valid_options else current_value
                st.session_state.inputs[field] = default_value
                return default_value
        st.session_state.inputs[field] = selected
        return selected

    # Render questions with unique keys and validation
    if current_category == "Personal Information":
        st.subheader("Personal Information")
        st.session_state.inputs['Age'] = st.number_input("Age", min_value=18, max_value=100, value=st.session_state.inputs['Age'], key="age_input")
        st.session_state.inputs['Gender'] = validated_selectbox("Gender", list(input_mappings['Gender'].keys()), st.session_state.inputs['Gender'], 'Gender', "gender_select")
        st.session_state.inputs['BMI'] = st.number_input("BMI", min_value=10.0, max_value=50.0, value=st.session_state.inputs['BMI'], step=0.1, key="bmi_input")

    elif current_category == "Lifestyle & Daily Habits":
        st.subheader("Lifestyle & Daily Habits")
        st.session_state.inputs['Smoking'] = validated_selectbox("Smoking", list(input_mappings['Smoking'].keys()), st.session_state.inputs['Smoking'], 'Smoking', "smoking_select")
        st.session_state.inputs['Alcohol_Use'] = validated_selectbox("Alcohol Use", list(input_mappings['Alcohol_Use'].keys()), st.session_state.inputs['Alcohol_Use'], 'Alcohol_Use', "alcohol_select")
        st.session_state.inputs['Sleep_Hours'] = st.number_input("Sleep Hours per Night", min_value=0, max_value=24, value=st.session_state.inputs['Sleep_Hours'], key="sleep_input")

    elif current_category == "Health Symptoms":
        st.subheader("Health Symptoms")
        st.session_state.inputs['Systolic_BP'] = st.number_input("Systolic Blood Pressure", min_value=50, max_value=250, value=st.session_state.inputs['Systolic_BP'], key="systolic_bp_input")
        st.session_state.inputs['Diastolic_BP'] = st.number_input("Diastolic Blood Pressure", min_value=30, max_value=150, value=st.session_state.inputs['Diastolic_BP'], key="diastolic_bp_input")
        st.session_state.inputs['Blood_Sugar'] = st.number_input("Blood Sugar (mg/dL)", min_value=50.0, max_value=300.0, value=st.session_state.inputs['Blood_Sugar'], step=0.1, key="blood_sugar_input")

    elif current_category == "Diet Type":
        st.subheader("Diet Type")
        st.write("How often do you eat fruits and vegetables in a week? (FruitVegFreq)")
        st.session_state.inputs['FruitVegFreq'] = validated_selectbox("FruitVegFreq", list(input_mappings['FruitVegFreq'].keys()), st.session_state.inputs['FruitVegFreq'], 'FruitVegFreq', "fruit_veg_select")
        st.write("Do you usually consume home-cooked meals or takeaways? (MealType)")
        st.session_state.inputs['MealType'] = validated_selectbox("MealType", list(input_mappings['MealType'].keys()), st.session_state.inputs['MealType'], 'MealType', "meal_type_select")
        st.write("How many times a week do you consume sugary snacks or beverages? (SugaryIntake)")
        st.session_state.inputs['SugaryIntake'] = validated_selectbox("SugaryIntake", list(input_mappings['SugaryIntake'].keys()), st.session_state.inputs['SugaryIntake'], 'SugaryIntake', "sugary_intake_select")
        st.write("Do you read nutrition labels before purchasing packaged foods? (ReadLabels)")
        st.session_state.inputs['ReadLabels'] = validated_selectbox("ReadLabels", list(input_mappings['ReadLabels'].keys()), st.session_state.inputs['ReadLabels'], 'ReadLabels', "read_labels_select")

    elif current_category == "Physical Activity":
        st.subheader("Physical Activity")
        st.write("How many days a week do you engage in any form of physical exercise? (ExerciseDays)")
        st.session_state.inputs['ExerciseDays'] = validated_selectbox("ExerciseDays", list(input_mappings['ExerciseDays'].keys()), st.session_state.inputs['ExerciseDays'], 'ExerciseDays', "exercise_days_select")
        st.write("Do you walk or cycle for at least 30 minutes a day? (WalkCycle)")
        st.session_state.inputs['WalkCycle'] = validated_selectbox("WalkCycle", list(input_mappings['WalkCycle'].keys()), st.session_state.inputs['WalkCycle'], 'WalkCycle', "walk_cycle_select")
        st.write("Do you have a desk job with little to no movement? (DeskJob)")
        st.session_state.inputs['DeskJob'] = validated_selectbox("DeskJob", list(input_mappings['DeskJob'].keys()), st.session_state.inputs['DeskJob'], 'DeskJob', "desk_job_select")
        st.write("Do you feel physically tired at the end of the day due to exertion? (PhysicalTired)")
        st.session_state.inputs['PhysicalTired'] = validated_selectbox("PhysicalTired", list(input_mappings['PhysicalTired'].keys()), st.session_state.inputs['PhysicalTired'], 'PhysicalTired', "physical_tired_select")

    elif current_category == "Stress Level":
        st.subheader("Stress Level")
        st.write("Do you often feel anxious or overwhelmed by daily tasks? (Anxious)")
        st.session_state.inputs['Anxious'] = validated_selectbox("Anxious", list(input_mappings['Anxious'].keys()), st.session_state.inputs['Anxious'], 'Anxious', "anxious_select")
        st.write("Do you get enough sleep without frequent interruptions? (GoodSleep)")
        st.session_state.inputs['GoodSleep'] = validated_selectbox("GoodSleep", list(input_mappings['GoodSleep'].keys()), st.session_state.inputs['GoodSleep'], 'GoodSleep', "good_sleep_select")
        st.write("Have you experienced mood swings or irritability recently? (MoodSwings)")
        st.session_state.inputs['MoodSwings'] = validated_selectbox("MoodSwings", list(input_mappings['MoodSwings'].keys()), st.session_state.inputs['MoodSwings'], 'MoodSwings', "mood_swings_select")
        st.write("Do you feel relaxed and focused during most of your day? (Relaxed)")
        st.session_state.inputs['Relaxed'] = validated_selectbox("Relaxed", list(input_mappings['Relaxed'].keys()), st.session_state.inputs['Relaxed'], 'Relaxed', "relaxed_select")

    elif current_category == "Symptoms":
        st.subheader("Symptoms")
        st.write("Do you often feel excessively thirsty or fatigued? (ThirstyFatigued)")
        st.session_state.inputs['ThirstyFatigued'] = validated_selectbox("ThirstyFatigued", list(input_mappings['ThirstyFatigued'].keys()), st.session_state.inputs['ThirstyFatigued'], 'ThirstyFatigued', "thirsty_fatigued_select")
        st.write("Have you noticed sudden weight loss or gain without trying? (WeightChange)")
        st.session_state.inputs['WeightChange'] = validated_selectbox("WeightChange", list(input_mappings['WeightChange'].keys()), st.session_state.inputs['WeightChange'], 'WeightChange', "weight_change_select")
        st.write("Do you experience chest pain, shortness of breath, or dizziness? (ChestPain)")
        st.session_state.inputs['ChestPain'] = validated_selectbox("ChestPain", list(input_mappings['ChestPain'].keys()), st.session_state.inputs['ChestPain'], 'ChestPain', "chest_pain_select")
        st.write("Have you had frequent urination recently? (FrequentUrination)")
        st.session_state.inputs['FrequentUrination'] = validated_selectbox("FrequentUrination", list(input_mappings['FrequentUrination'].keys()), st.session_state.inputs['FrequentUrination'], 'FrequentUrination', "frequent_urination_select")
        st.write("Have you had blurred vision recently? (BlurredVision)")
        st.session_state.inputs['BlurredVision'] = validated_selectbox("BlurredVision", list(input_mappings['BlurredVision'].keys()), st.session_state.inputs['BlurredVision'], 'BlurredVision', "blurred_vision_select")

    elif current_category == "Family History":
        st.subheader("Family History")
        st.write("Has any close family member (parents, siblings) been diagnosed with diabetes? (FamDiabetes)")
        st.session_state.inputs['FamDiabetes'] = validated_selectbox("FamDiabetes", list(input_mappings['FamDiabetes'].keys()), st.session_state.inputs['FamDiabetes'], 'FamDiabetes', "fam_diabetes_select")
        st.write("Is there a history of heart disease in your immediate family (before age 60)? (FamHeartDisease)")
        st.session_state.inputs['FamHeartDisease'] = validated_selectbox("FamHeartDisease", list(input_mappings['FamHeartDisease'].keys()), st.session_state.inputs['FamHeartDisease'], 'FamHeartDisease', "fam_heart_disease_select")
        st.write("Have any of your blood relatives been diagnosed with high blood pressure or hypertension? (FamHypertension)")
        st.session_state.inputs['FamHypertension'] = validated_selectbox("FamHypertension", list(input_mappings['FamHypertension'].keys()), st.session_state.inputs['FamHypertension'], 'FamHypertension', "fam_hypertension_select")
        st.write("Has anyone in your family required long-term medication for cholesterol issues? (FamCholesterol)")
        st.session_state.inputs['FamCholesterol'] = validated_selectbox("FamCholesterol", list(input_mappings['FamCholesterol'].keys()), st.session_state.inputs['FamCholesterol'], 'FamCholesterol', "fam_cholesterol_select")
        st.write("Do your parents or grandparents have a history of strokes or cardiovascular events? (FamStroke)")
        st.session_state.inputs['FamStroke'] = validated_selectbox("FamStroke", list(input_mappings['FamStroke'].keys()), st.session_state.inputs['FamStroke'], 'FamStroke', "fam_stroke_select")
        st.write("Is there a known pattern of obesity in your family members? (FamObesity)")
        st.session_state.inputs['FamObesity'] = validated_selectbox("FamObesity", list(input_mappings['FamObesity'].keys()), st.session_state.inputs['FamObesity'], 'FamObesity', "fam_obesity_select")
        st.write("Have any close relatives undergone surgeries related to heart disease or blocked arteries? (FamHeartSurgery)")
        st.session_state.inputs['FamHeartSurgery'] = validated_selectbox("FamHeartSurgery", list(input_mappings['FamHeartSurgery'].keys()), st.session_state.inputs['FamHeartSurgery'], 'FamHeartSurgery', "fam_heart_surgery_select")
        st.write("Have any of your siblings or parents faced complications due to chronic illnesses? (FamChronic)")
        st.session_state.inputs['FamChronic'] = validated_selectbox("FamChronic", list(input_mappings['FamChronic'].keys()), st.session_state.inputs['FamChronic'], 'FamChronic', "fam_chronic_select")
        st.write("Is there a history of kidney disease or dialysis among your family members? (FamKidney)")
        st.session_state.inputs['FamKidney'] = validated_selectbox("FamKidney", list(input_mappings['FamKidney'].keys()), st.session_state.inputs['FamKidney'], 'FamKidney', "fam_kidney_select")
        st.write("Have any of your first-degree relatives passed away due to chronic health issues before age 65? (FamEarlyDeath)")
        st.session_state.inputs['FamEarlyDeath'] = validated_selectbox("FamEarlyDeath", list(input_mappings['FamEarlyDeath'].keys()), st.session_state.inputs['FamEarlyDeath'], 'FamEarlyDeath', "fam_early_death_select")

    # Navigation buttons
    st.markdown("<br>", unsafe_allow_html=True)
    with st.container():
        st.markdown('<div class="button-container">', unsafe_allow_html=True)
        col1, col2, col3, col4 = st.columns([1, 1, 1, 1])
        
        # Validate inputs for the current category
        is_valid = validate_category_inputs(current_category, st.session_state.inputs)
        
        with col2:
            if st.button("Prev", key=f"prev_button_{current_category}", disabled=(st.session_state.category_index == 0)):
                st.session_state.category_index = max(0, st.session_state.category_index - 1)
        
        with col3:
            if current_category != "Family History":
                if st.button("Next", key=f"next_button_{current_category}", disabled=(not is_valid or st.session_state.category_index == len(categories) - 1)):
                    st.session_state.category_index = min(len(categories) - 1, st.session_state.category_index + 1)
            else:
                if st.button("🔮Predict Disease", key="predict_button", disabled=not is_valid):
                    try:
                        # Collect all inputs
                        input_data = st.session_state.inputs.copy()

                        # Encode categorical inputs using original dataset values
                        for feature in input_data:
                            if feature in encoders:
                                value = input_data[feature]
                                if feature in input_mappings:
                                    value = input_mappings[feature][value]
                                try:
                                    encoded = encoders[feature].transform([value])[0]
                                    input_data[feature] = encoded
                                except ValueError as e:
                                    st.error(f"Encoding error for {feature}: {value}. Expected one of {list(encoders[feature].classes_)}. Please check your input.")
                                    st.stop()

                        # Predict for each target
                        predictions = {}
                        risk_level_map = {0: "Low", 1: "Medium", 2: "High"}
                        risk_level_map_reverse = {"Low": 0, "Medium": 1, "High": 2}

                        for target, model in [
                            ('Risk_Level_Diabetes', diabetes_model),
                            ('Risk_Level_Hypertension', hypertension_model),
                            ('Risk_Level_Heart_Disease', heart_disease_model)
                        ]:
                            expected_features = feature_sets[target]
                            # Filter input_data to include only expected features
                            filtered_input_data = {k: v for k, v in input_data.items() if k in expected_features}
                            missing_features = [f for f in expected_features if f not in filtered_input_data]
                            if missing_features:
                                st.error(f"Missing features for {target}: {missing_features}. Cannot proceed with prediction.")
                                st.stop()

                            # Create DataFrame with correct feature order
                            input_df = pd.DataFrame([filtered_input_data])[expected_features]
                            
                            # Make prediction
                            try:
                                pred_raw = model.predict(input_df)
                                prob = model.predict_proba(input_df)[0]
                            except Exception as e:
                                st.error(f"Model prediction failed for {target}: {str(e)}. Please ensure the model is compatible with the input data.")
                                st.stop()

                            # Handle pred_raw
                            if isinstance(pred_raw, (list, np.ndarray)):
                                pred_value = pred_raw[0]
                            else:
                                pred_value = pred_raw

                            # Convert prediction to integer index
                            try:
                                if isinstance(pred_value, str):
                                    if pred_value not in risk_level_map_reverse:
                                        st.error(f"Invalid prediction label for {target}: {pred_value}. Expected 'Low', 'Medium', or 'High'.")
                                        st.stop()
                                    pred = risk_level_map_reverse[pred_value]
                                else:
                                    pred = int(float(pred_value))
                            except (ValueError, TypeError) as e:
                                st.error(f"Cannot process prediction for {target}: {pred_raw}. Expected a class label ('Low', 'Medium', 'High') or integer (0, 1, 2).")
                                st.stop()

                            # Validate pred is a valid index for prob
                            if not isinstance(pred, int) or pred < 0 or pred >= len(prob):
                                st.error(f"Invalid prediction index for {target}: {pred}. Must be an integer between 0 and {len(prob)-1}.")
                                st.stop()

                            # Validate probability array
                            if not isinstance(prob, np.ndarray) or len(prob) < 1 or np.any(np.isnan(prob)):
                                st.error(f"Invalid probability output for {target}: {prob}. Expected a non-empty probability array without NaN.")
                                st.stop()

                            # Calculate severity as the probability of the predicted class
                            severity = int(prob[pred] * 100)

                            predictions[target] = {'pred': pred, 'severity': severity, 'prob': prob}

                        # Map predictions to risk levels
                        st.session_state.diabetes_risk = risk_level_map[predictions['Risk_Level_Diabetes']['pred']]
                        st.session_state.hypertension_risk = risk_level_map[predictions['Risk_Level_Hypertension']['pred']]
                        st.session_state.heart_disease_risk = risk_level_map[predictions['Risk_Level_Heart_Disease']['pred']]
                        st.session_state.predictions = predictions
                        st.session_state.show_results = True

                        # Save predictions to predictions.csv
                        patient_id = f"PT-{int(datetime.now().timestamp())}"  # Generate unique Patient ID
                        save_prediction(patient_id, "Diabetes", st.session_state.diabetes_risk)
                        save_prediction(patient_id, "Hypertension", st.session_state.hypertension_risk)
                        save_prediction(patient_id, "Heart Disease", st.session_state.heart_disease_risk)

                        # Save inputs and predictions to Word document
                        save_to_word(
                            patient_id,
                            st.session_state.inputs,
                            st.session_state.diabetes_risk,
                            st.session_state.hypertension_risk,
                            st.session_state.heart_disease_risk
                        )

                    except Exception as e:
                        st.error(f"Unexpected error during prediction: {str(e)}. Please check your inputs and model files.")
                        st.stop()
        
        with col4:
            if current_category == "Family History":
                if st.button("Next", key=f"next_button_{current_category}", disabled=(not is_valid or st.session_state.category_index == len(categories) - 1)):
                    st.session_state.category_index = min(len(categories) - 1, st.session_state.category_index + 1)
        
        # Display validation message if inputs are incomplete
        if not is_valid:
            st.warning("Please answer all questions in this category with valid inputs before proceeding.")
        
        st.markdown('</div>', unsafe_allow_html=True)

    # Render prediction results using Streamlit components
    if st.session_state.show_results and st.session_state.predictions:
        with st.container():
            st.markdown(
                """
                <div class="prediction-container">
                    <h2 style='text-align: center; color: #000080;'>Prediction Results</h2>
                </div>
                """,
                unsafe_allow_html=True
            )

            def get_recommendations(risk, disease):
                if risk == "Low":
                    if disease == "Diabetes":
                        return [
                            "Diet: Maintain low-sugar, high-fiber diet (e.g., whole grains, vegetables).",
                            "Activity: Continue 30 mins/day moderate exercise (e.g., walking).",
                            "Monitoring: Check blood sugar annually."
                        ]
                    elif disease == "Heart Disease":
                        return [
                            "Diet: Keep low-fat, heart-healthy diet (e.g., omega-3 rich foods).",
                            "Activity: Maintain 150 mins/week aerobic exercise.",
                            "Monitoring: Annual cholesterol and BP checks."
                        ]
                    else:  # Hypertension
                        return [
                            "Diet: Reduce sodium, increase potassium-rich foods (e.g., bananas).",
                            "Activity: Continue daily physical activity.",
                            "Monitoring: Regular BP checks."
                        ]
                elif risk == "Medium":
                    if disease == "Diabetes":
                        return [
                            "Diet: Reduce refined carbs, limit sugary drinks.",
                            "Activity: Increase to 45 mins/day exercise (e.g., brisk walking).",
                            "Medical: Consult for A1C test every 6 months."
                        ]
                    elif disease == "Heart Disease":
                        return [
                            "Diet: Avoid trans fats, limit red meat.",
                            "Activity: 45 mins/day cardio (e.g., cycling).",
                            "Medical: ECG and lipid profile every 6 months."
                        ]
                    else:  # Hypertension
                        return [
                            "Diet: Low-sodium diet, avoid processed foods.",
                            "Activity: 45 mins/day exercise, include yoga.",
                            "Medical: Monitor BP weekly, consult doctor."
                        ]
                else:  # High
                    if disease == "Diabetes":
                        return [
                            "Diet: Strict low-carb, low-sugar diet.",
                            "Activity: 1 hr/day structured exercise, consult trainer.",
                            "Medical: Urgent A1C test, endocrinologist visit."
                        ]
                    elif disease == "Heart Disease":
                        return [
                            "Diet: Low-fat, low-cholesterol diet.",
                            "Activity: 1 hr/day cardio, under medical supervision.",
                            "Medical: Urgent cardiologist visit, ECG, stress test."
                        ]
                    else:  # Hypertension
                        return [
                            "Diet: No added salt, high-fiber diet.",
                            "Activity: 1 hr/day exercise, avoid high stress.",
                            "Medical: Urgent BP evaluation, medication review."
                        ]

            # Display predictions in a row
            st.markdown('<div class="prediction-row">', unsafe_allow_html=True)
            cols = st.columns(3)
            for idx, (target, disease, risk) in enumerate([
                ('Risk_Level_Diabetes', 'Diabetes', st.session_state.diabetes_risk),
                ('Risk_Level_Hypertension', 'Hypertension', st.session_state.hypertension_risk),
                ('Risk_Level_Heart_Disease', 'Heart Disease', st.session_state.heart_disease_risk)
            ]):
                with cols[idx]:
                    st.markdown(
                        f"""
                        <div class="prediction-card">
                            <h3>{disease}</h3>
                            <p><strong>Risk Level:</strong> {risk}</p>
                            <div class="recommendation">
                                <strong>Recommendations:</strong><br>
                                {"<br>".join([f"• {rec}" for rec in get_recommendations(risk, disease)])}
                            </div>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
            
            # Add visualization chart
            st.markdown(
                """
                <div class="chart-header">
                    <h2>Risk Level Visualization</h2>
                </div>
                """,
                unsafe_allow_html=True
            )
            
            # Prepare data for the chart
            diseases = ['Diabetes', 'Hypertension', 'Heart Disease']
            risk_levels = [st.session_state.diabetes_risk, st.session_state.hypertension_risk, st.session_state.heart_disease_risk]
            
            # Convert risk levels to numerical values for the chart
            risk_values = []
            for risk in risk_levels:
                if risk == "Low":
                    risk_values.append(1)
                elif risk == "Medium":
                    risk_values.append(2)
                else:  # High
                    risk_values.append(3)
            
            # Create the bar chart with adjusted size
            fig, ax = plt.subplots(figsize=(7, 3))
            
            # Blue color gradient: light blue for Low, medium blue for Medium, navy blue for High
            colors = ['#ADD8E6', '#1E90FF', '#000080']
            bars = ax.bar(diseases, risk_values, color=colors, width=0.6)
            
            # Customize the chart
            ax.set_ylim(0, 3.5)
            ax.set_yticks([1, 2, 3])
            ax.set_yticklabels(['Low', 'Medium', 'High'])
            ax.set_ylabel('Risk Level', fontsize=10)
            ax.set_xlabel('Diseases', fontsize=10)
            ax.set_title('Disease Risk Levels', fontsize=12)
            
            # Add value labels on top of each bar with adjusted font size
            for bar, risk in zip(bars, risk_levels):
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height,
                        f'{risk}',
                        ha='center', va='bottom', fontsize=10)
            
            # Set white background and remove top/right spines
            fig.patch.set_facecolor('white')
            ax.set_facecolor('white')
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
            
            # Adjust layout to prevent label cutoff
            plt.tight_layout()
            
            # Display the chart
            st.pyplot(fig)
            st.markdown('</div>', unsafe_allow_html=True)

            # Center the Close Results button
            with st.container():
                st.markdown('<div class="center-button-container">', unsafe_allow_html=True)
                if st.button("Close Results", key="close_results_button"):
                    st.session_state.show_results = False
                    st.session_state.predictions = None
                    st.session_state.category_index = 0  # Redirect to Personal Information
                st.markdown('</div>', unsafe_allow_html=True)